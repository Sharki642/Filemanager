#importing the tkinter module
import tkinter as tk
from docx import Document

print("hello world!")
#starting the window
root = tk.Tk()

def create_protokoll_de():
    print("Creating German Protokoll...")
    window = tk.Toplevel(root)
    window.title("German Protokoll")
    window.minsize(300, 200)
    tk.Label(window, text="Dateiname:").pack()
    textfeld_ger = tk.Text(window, height=2, width=30)
    textfeld_ger.pack()
    tk.Label(window, text="Überschrift:").pack()
    textfeld_ueberschrift = tk.Text(window, height=2, width=30)
    textfeld_ueberschrift.pack()

    def erstelle_date_ger():
        eingabe_name = textfeld_ger.get("1.0", tk.END).strip()
        eingabe_ueberschrift = textfeld_ueberschrift.get("1.0", tk.END).strip()
        if not eingabe_name:
            tk.Label(window, text="Bitte einen Dateinamen eingeben!", fg="red").pack()
            return
        if not eingabe_ueberschrift:
            tk.Label(window, text="Bitte eine Überschrift eingeben!", fg="red").pack()
            return
        doc = Document("vorlage_deutsch.docx")
        alle_paragraphen = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        alle_paragraphen.append(para)
        for para in alle_paragraphen:
            if "<Überschrift>" in para.text:
                voller_text = para.text.replace("<Überschrift>", eingabe_ueberschrift)
                for i, run in enumerate(para.runs):
                    run.text = voller_text if i == 0 else ""
        doc.save(f"{eingabe_name}.docx")
        tk.Label(window, text=f"Datei '{eingabe_name}.docx' erstellt!", fg="green").pack()

    tk.Button(window, text="Datei erstellen", command=erstelle_date_ger).pack()

def create_protokoll_en():
    print("Creating English Protokoll...")
    window = tk.Toplevel(root)
    window.title("English Protokoll")
    window.minsize(300, 200)
    tk.Label(window, text="File name:").pack()
    textfeld_eng = tk.Text(window, height=2, width=30)
    textfeld_eng.pack()
    tk.Label(window, text="Title:").pack()
    textfeld_title = tk.Text(window, height=2, width=30)
    textfeld_title.pack()

    def erstelle_date_eng():
        eingabe_name = textfeld_eng.get("1.0", tk.END).strip()
        eingabe_title = textfeld_title.get("1.0", tk.END).strip()
        if not eingabe_name:
            tk.Label(window, text="Please enter a file name!", fg="red").pack()
            return
        if not eingabe_title:
            tk.Label(window, text="Please enter a title!", fg="red").pack()
            return
        doc = Document("vorlage_english.docx")
        alle_paragraphen = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        alle_paragraphen.append(para)
        for para in alle_paragraphen:
            if "<name of exercise>" in para.text:
                voller_text = para.text.replace("<name of exercise>", eingabe_title)
                for i, run in enumerate(para.runs):
                    run.text = voller_text if i == 0 else ""
        doc.save(f"{eingabe_name}.docx")
        tk.Label(window, text=f"File '{eingabe_name}.docx' created!", fg="green").pack()

    tk.Button(window, text="Create File", command=erstelle_date_eng).pack()

#title
root.title("FileManager")
root.minsize(400, 400)
#inside

root.configure(bg="White")
tk.Label(root, text="Welcome to FileManager").pack()
image = tk.PhotoImage(file="image3.png").subsample(2, 2)
tk.Label(root, image=image).pack()


Protokoll_de = tk.Button(root, text="Create German Protokoll", command=create_protokoll_de)
Protokoll_de.pack()

Protokoll_en = tk.Button(root, text="Create English Protokoll", command=create_protokoll_en)
Protokoll_en.pack()



root.mainloop()